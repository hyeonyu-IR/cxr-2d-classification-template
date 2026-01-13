#!/usr/bin/env python
"""
scripts/04_make_manuscript_doc.py

Generate a manuscript-ready Word document summarizing a completed run.
"""

from pathlib import Path
import argparse
import json
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
import matplotlib.pyplot as plt
import numpy as np


def load_json(p: Path):
    with p.open("r", encoding="utf-8") as f:
        return json.load(f)


def resolve_run_dir(run: str, runs_root: Path) -> Path:
    if run == "latest":
        meta = load_json(runs_root / "_latest_run.json")
        return Path(meta["run_dir"]).resolve()
    return Path(run).resolve()


def add_dataframe(doc: Document, df: pd.DataFrame, title: str):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


def add_images_from_folder(doc: Document, folder: Path, title: str):
    if not folder.exists():
        return

    doc.add_heading(title, level=2)

    exts = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif", ".webp"}

    # Only include files (not directories like .ipynb_checkpoints)
    images = [p for p in sorted(folder.iterdir())
              if p.is_file() and p.suffix.lower() in exts]

    if not images:
        doc.add_paragraph("(No image files found.)")
        return

    for img in images:
        doc.add_paragraph(img.name)
        try:
            doc.add_picture(str(img), width=Inches(6))
        except Exception as e:
            # Do not fail the entire document build because of one file
            doc.add_paragraph(f"[Skipped: could not embed image: {img.name} | {type(e).__name__}: {e}]")
        doc.add_page_break()


def save_training_curves(history_csv: Path, out_dir: Path) -> list[Path]:
    """
    Create training curve PNG(s) from history.csv.
    Matches this repo's history schema (epoch, stage, train_loss, train_acc_quick, val_loss, val_acc, val_ap, test_ap).
    Returns a list of created PNG paths.
    """
    if not history_csv.exists():
        return []

    out_dir.mkdir(parents=True, exist_ok=True)
    df = pd.read_csv(history_csv)
    df.columns = [c.strip() for c in df.columns]

    x = df["epoch"] if "epoch" in df.columns else (np.arange(len(df)) + 1)

    # Optional: mark the stage transition (head -> finetune) if stage column exists
    stage_change_epoch = None
    if "stage" in df.columns:
        stages = df["stage"].astype(str).values
        for i in range(1, len(stages)):
            if stages[i] != stages[i - 1]:
                stage_change_epoch = int(x.iloc[i]) if hasattr(x, "iloc") else int(x[i])
                break

    created: list[Path] = []

    def _finalize_plot(ax, title: str, out_png: Path):
        ax.set_xlabel("Epoch")
        ax.set_title(title)
        ax.grid(True, alpha=0.3)
        ax.legend()

        # Add a vertical marker for stage change if detected
        if stage_change_epoch is not None:
            ax.axvline(stage_change_epoch, linestyle="--", linewidth=1)
            ax.text(stage_change_epoch, ax.get_ylim()[1], " stage change", va="top")

        plt.tight_layout()
        plt.savefig(out_png, bbox_inches="tight")
        plt.close()

    # 1) Loss curve: train_loss vs val_loss
    if "train_loss" in df.columns or "val_loss" in df.columns:
        fig, ax = plt.subplots(figsize=(7.0, 4.0), dpi=200)
        if "train_loss" in df.columns:
            ax.plot(x, df["train_loss"].values, label="train_loss")
        if "val_loss" in df.columns:
            ax.plot(x, df["val_loss"].values, label="val_loss")
        ax.set_ylabel("Loss")
        out_png = out_dir / "training_curve_loss.png"
        _finalize_plot(ax, "Training Curve: Loss", out_png)
        created.append(out_png)

    # 2) Accuracy curve: train_acc_quick vs val_acc
    if "train_acc_quick" in df.columns or "val_acc" in df.columns:
        fig, ax = plt.subplots(figsize=(7.0, 4.0), dpi=200)
        if "train_acc_quick" in df.columns:
            ax.plot(x, df["train_acc_quick"].values, label="train_acc_quick")
        if "val_acc" in df.columns:
            ax.plot(x, df["val_acc"].values, label="val_acc")
        ax.set_ylabel("Accuracy")
        out_png = out_dir / "training_curve_accuracy.png"
        _finalize_plot(ax, "Training Curve: Accuracy", out_png)
        created.append(out_png)

    # 3) AP curve: val_ap and (optional) test_ap
    if "val_ap" in df.columns or "test_ap" in df.columns:
        fig, ax = plt.subplots(figsize=(7.0, 4.0), dpi=200)
        if "val_ap" in df.columns:
            ax.plot(x, df["val_ap"].values, label="val_ap")
        if "test_ap" in df.columns:
            ax.plot(x, df["test_ap"].values, label="test_ap")
        ax.set_ylabel("Average Precision")
        out_png = out_dir / "training_curve_ap.png"
        _finalize_plot(ax, "Training Curve: Average Precision", out_png)
        created.append(out_png)

    return created


def main():
    parser = argparse.ArgumentParser(description="Generate manuscript-ready Word document.")
    parser.add_argument("--run", default="latest", help="Run directory or 'latest'")
    parser.add_argument("--runs-root", default="outputs/runs")
    parser.add_argument("--out-name", default="manuscript_report.docx")
    args = parser.parse_args()

    run_dir = resolve_run_dir(args.run, Path(args.runs_root))
    out_path = run_dir / args.out_name

    doc = Document()
    doc.add_heading("Model Training and Evaluation Report", 0)
    doc.add_paragraph(f"Run directory: {run_dir}")
    doc.add_paragraph(f"Generated: {datetime.now().isoformat()}")

    # Environment
    env_path = run_dir / "env_report.json"
    if env_path.exists():
        doc.add_heading("Environment", level=1)
        env = load_json(env_path)
        for k, v in env.items():
            doc.add_paragraph(f"{k}: {v}")

    # Config
    cfg_path = run_dir / "config.json"
    if cfg_path.exists():
        doc.add_heading("Configuration", level=1)
        cfg = load_json(cfg_path)
        for k, v in cfg.items():
            doc.add_paragraph(f"{k}: {v}")
    
    
    # History + training curves
    hist_path = run_dir / "history.csv"
    
    if hist_path.exists():
        doc.add_heading("Training", level=1)
    
        fig_dir = run_dir / "report" / "figures"
        curve_paths = save_training_curves(hist_path, fig_dir)
    
        if curve_paths:
            doc.add_heading("Training Curves", level=2)
            for p in curve_paths:
                doc.add_paragraph(p.name)
                doc.add_picture(str(p), width=Inches(6))
            doc.add_paragraph("")
        else:
            doc.add_paragraph(
                "Training curves could not be generated "
                "(missing expected columns in history.csv)."
            )
    
        # Training history table
        df = pd.read_csv(hist_path)
        add_dataframe(doc, df, "Training History (per epoch)")


    # Then include the table
    df = pd.read_csv(hist_path)
    add_dataframe(doc, df, "Training History (per epoch)")

    
    # Evaluation summary
    eval_path = run_dir / "eval_summary.json"
    if eval_path.exists():
        doc.add_heading("Evaluation Summary", level=1)
        eval_summary = load_json(eval_path)
        for k, v in eval_summary.items():
            doc.add_paragraph(f"{k}: {v}")

    
    # Error tables
    for err_csv in run_dir.glob("errors_*.csv"):
        df = pd.read_csv(err_csv)
        add_dataframe(doc, df, f"Errors: {err_csv.name}")

    
    # Grad-CAM images
    add_images_from_folder(doc, run_dir / "gradcam" / "FP", "Grad-CAM: False Positives")
    add_images_from_folder(doc, run_dir / "gradcam" / "FN", "Grad-CAM: False Negatives")

    doc.save(out_path)
    print(f"Manuscript document written to: {out_path}")


if __name__ == "__main__":
    main()
